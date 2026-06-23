"""
HR News Digest – Standalone mit Textverarbeitung (ohne KI)
==========================================================
Scrapt deutsche HR-Quellen, holt die vollen Artikel-Texte,
und erstellt automatisch eine echte Zusammenfassung pro Artikel.

Nutzung:
    pip install feedparser trafilatura python-docx
    python hr_digest.py
"""

import feedparser
import trafilatura
import re
from datetime import datetime, timedelta, timezone
from time import mktime
from collections import defaultdict
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Konfiguration ──────────────────────────────────────────

FEEDS = [
    # ── Frei zugänglich (kein Paywall) ──
    ("Haufe Personal",        "https://www.haufe.de/personal/rss_366_488.xml"),
    ("HR Journal",            "https://www.hrjournal.de/feed/"),
    ("Personalwirtschaft",    "https://www.personalwirtschaft.de/rss/feed.xml"),
    ("Persoblogger",          "https://persoblogger.de/feed/"),
    ("HRM Online",            "https://www.hrm.de/feed/"),
    ("Arbeitsrecht.de",       "https://www.arbeitsrecht.de/feed/"),
    ("t3n Karriere",          "https://t3n.de/rss.xml"),  # wird per Keyword gefiltert
    # ── Paywall-Quellen (RSS-Vorschau oft ausreichend) ──
    ("FAZ Beruf & Chance",    "https://www.faz.net/rss/aktuell/beruf-chance/"),
    ("Handelsblatt Karriere", "https://www.handelsblatt.com/contentexport/feed/karriere"),
    ("Spiegel Karriere",      "https://www.spiegel.de/karriere/index.rss"),
]

CATEGORIES = {
    "⚖️ Arbeitsrecht & Regulierung": [
        "arbeitsrecht", "kündigungsschutz", "kündigung", "tarifvertrag",
        "betriebsrat", "arbeitszeitgesetz", "befristung", "gesetz",
        "regulierung", "richtlinie", "compliance", "datenschutz",
    ],
    "🔍 Recruiting & Arbeitsmarkt": [
        "recruiting", "fachkräftemangel", "bewerbung", "arbeitsmarkt",
        "stellenanzeige", "employer branding", "talent", "einstellung",
    ],
    "📈 Personalentwicklung & Weiterbildung": [
        "weiterbildung", "personalentwicklung", "onboarding", "training",
        "qualifizierung", "kompetenz", "skill", "fortbildung", "coaching",
    ],
    "💰 Vergütung & Benefits": [
        "gehalt", "gehaltserhöhung", "vergütung", "mindestlohn", "bonus",
        "lohnabrechnung", "payroll", "sozialversicherung", "benefit",
        "betriebsrente", "lohn", "entgelt",
    ],
    "🏠 New Work & Arbeitskultur": [
        "homeoffice", "remote work", "hybrides arbeiten", "new work",
        "work-life-balance", "unternehmenskultur", "führung", "leadership",
        "change management", "agil", "diversity", "inklusion",
    ],
    "❤️ Gesundheit & Wellbeing": [
        "mental health", "gesundheit", "bgm", "burnout", "stress",
        "betriebliche gesundheit", "wellbeing", "prävention",
    ],
    "🤖 Digitalisierung im HR": [
        "hr-software", "digitalisierung", "ki ", "künstliche intelligenz",
        "automatisierung", "hr-tech", "people analytics", "digital",
    ],
}

ALL_KEYWORDS = []
for kws in CATEGORIES.values():
    ALL_KEYWORDS.extend(kws)
ALL_KEYWORDS += ["personal", "hr", "human resources", "mitarbeiter",
                  "personalplanung", "fluktuation", "elternzeit"]

MAX_AGE_DAYS = 7
MAX_ARTICLES = 30

# ── Hilfsfunktionen ────────────────────────────────────────

def clean_html(text):
    return re.sub(r"<[^>]+>", "", text or "").strip()


def parse_date(entry):
    for attr in ("published_parsed", "updated_parsed"):
        val = getattr(entry, attr, None)
        if val:
            try:
                return datetime.fromtimestamp(mktime(val), tz=timezone.utc)
            except (ValueError, OverflowError):
                continue
    return None


def matches_any(text, keywords):
    t = text.lower()
    return any(k in t for k in keywords)


def categorize(text):
    text = text.lower()
    best = "📌 Sonstiges"
    best_score = 0
    for cat, keywords in CATEGORIES.items():
        score = sum(1 for k in keywords if k in text)
        if score > best_score:
            best_score = score
            best = cat
    return best


# ── Paywall-Erkennung ──────────────────────────────────────

PAYWALL_DOMAINS = {"faz.net", "handelsblatt.com", "spiegel.de", "zeit.de",
                   "sueddeutsche.de", "welt.de", "wiwo.de", "manager-magazin.de"}

PAYWALL_SIGNALS = [
    "paywall", "premium-content", "plus-content", "abo-content",
    "regwall", "piano-offer", "subscribe", "registrieren sie sich",
    "lesen sie den vollständigen artikel", "jetzt freischalten",
    "exklusiv für abonnenten", "s+ artikel",
]

def is_paywall_domain(url):
    """Prüft ob die URL zu einer bekannten Paywall-Seite gehört."""
    return any(d in url.lower() for d in PAYWALL_DOMAINS)

def looks_like_paywall(text):
    """Erkennt ob der extrahierte Text nach Paywall-Fragment aussieht."""
    if not text:
        return True
    # Sehr kurzer Text = wahrscheinlich nur Anriss
    if len(text.split()) < 50:
        return True
    text_lower = text.lower()
    hits = sum(1 for s in PAYWALL_SIGNALS if s in text_lower)
    return hits >= 2


# ── Artikel-Text holen (Multi-Strategie) ──────────────────

def extract_rss_fulltext(entry):
    """Holt den maximal verfügbaren Text aus dem RSS-Feed selbst.
    
    Viele Feeds liefern in 'content:encoded' oder 'content' den
    vollständigen Artikeltext – das ist kein Paywall-Umgehen,
    sondern das Lesen dessen was der Verlag im Feed bereitstellt.
    """
    texts = []

    # content:encoded (oft Volltext bei Blogs/Fachmagazinen)
    for content in getattr(entry, "content", []):
        val = content.get("value", "")
        if val:
            texts.append(clean_html(val))

    # Standard-Summary
    summary = clean_html(getattr(entry, "summary", ""))
    if summary:
        texts.append(summary)

    # Description (manchmal anders als Summary)
    desc = clean_html(getattr(entry, "description", ""))
    if desc and desc != summary:
        texts.append(desc)

    if not texts:
        return None

    # Den längsten Text nehmen
    best = max(texts, key=len)
    return best if len(best) > 80 else None


def fetch_article_text(url):
    """Lädt die Webseite und extrahiert den Haupttext."""
    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            return None
        text = trafilatura.extract(
            downloaded,
            include_comments=False,
            include_tables=False,
            deduplicate=True,
            favor_precision=False,   # Lieber mehr Text als zu wenig
            include_links=False,
        )
        return text
    except Exception as e:
        print(f"      ⚠️  Text-Extraktion fehlgeschlagen: {e}")
        return None


def get_best_text(entry, url):
    """Holt den besten verfügbaren Text über mehrere Strategien.
    
    Reihenfolge:
    1. Volltext von der Webseite (trafilatura)
    2. Falls Paywall erkannt → RSS-Volltext als Fallback
    3. Falls beides mager → kombiniere was da ist
    """
    rss_text = extract_rss_fulltext(entry)
    web_text = None
    source = "rss"

    # Bei bekannten Paywall-Domains: RSS zuerst, Web nur als Bonus
    if is_paywall_domain(url):
        print(f"      🔒 Paywall-Quelle erkannt, nutze RSS-Inhalt")
        web_text = fetch_article_text(url)
        if web_text and not looks_like_paywall(web_text):
            source = "web"
        elif rss_text and len(rss_text) > 100:
            source = "rss"
            web_text = None  # RSS ist besser
        else:
            source = "partial"
    else:
        web_text = fetch_article_text(url)
        if web_text and len(web_text) > 100:
            source = "web"

    # Bestes Ergebnis auswählen
    if source == "web" and web_text:
        return web_text, source
    elif rss_text and len(rss_text) > len(web_text or ""):
        return rss_text, "rss"
    elif web_text:
        return web_text, "partial"
    elif rss_text:
        return rss_text, "rss"
    else:
        return None, "none"


# ── Extraktive Zusammenfassung ─────────────────────────────

STOPWORDS = {
    "der", "die", "das", "und", "oder", "für", "von", "mit",
    "ist", "sind", "ein", "eine", "auf", "bei", "nach", "wie",
    "sich", "auch", "den", "dem", "des", "im", "in", "zu",
    "nicht", "als", "an", "es", "hat", "wird", "zum", "zur",
    "über", "aus", "vor", "noch", "nur", "mehr", "so", "aber",
    "was", "wenn", "man", "kann", "haben", "werden", "dass",
    "einem", "einer", "diese", "dieser", "dieses", "bis",
}


def split_sentences(text):
    """Teilt Text in Sätze (einfacher deutscher Satz-Splitter)."""
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ])', text)
    return [s.strip() for s in sentences if len(s.split()) >= 8]


def score_sentence(sentence, title_words):
    """Bewertet einen Satz nach Relevanz."""
    words = sentence.lower().split()
    score = 0.0

    # Bonus: Wörter aus dem Titel
    for w in words:
        if w in title_words:
            score += 2.0

    # Bonus: HR-Keywords
    sent_lower = sentence.lower()
    for kw in ALL_KEYWORDS:
        if kw in sent_lower:
            score += 1.5

    # Bonus: Zahlen und Fakten
    if re.search(r'\d+[,.]?\d*\s*(%|Prozent|Euro|Milliard|Million)', sentence):
        score += 2.0

    # Malus: Sehr lange Sätze
    if len(words) > 40:
        score *= 0.7

    # Bonus: Signalwörter am Satzanfang
    starters = ["demnach", "laut", "insgesamt", "erstmals", "künftig",
                 "besonders", "entscheidend", "wichtig", "zentral",
                 "das ergebnis", "die studie", "experten"]
    for s in starters:
        if sent_lower.startswith(s):
            score += 1.0
            break

    return score


def summarize_text(text, title, num_sentences=3):
    """Extraktive Zusammenfassung: Wählt die relevantesten Sätze."""
    if not text or len(text) < 100:
        return None

    sentences = split_sentences(text)
    if not sentences:
        return None

    title_words = {
        w.lower() for w in title.split()
        if w.lower() not in STOPWORDS and len(w) > 2
    }

    scored = []
    for i, sent in enumerate(sentences):
        score = score_sentence(sent, title_words)
        # Positionsbonus (Lead-Prinzip bei Nachrichtenartikeln)
        if i == 0:
            score += 3.0
        elif i == 1:
            score += 1.5
        elif i == 2:
            score += 0.5
        scored.append((score, i, sent))

    scored.sort(key=lambda x: x[0], reverse=True)
    top = scored[:num_sentences]
    top.sort(key=lambda x: x[1])  # Zurück in Originalreihenfolge

    return " ".join(s[2] for s in top)


# ── RSS Scraping ───────────────────────────────────────────

def scrape():
    cutoff = datetime.now(timezone.utc) - timedelta(days=MAX_AGE_DAYS)
    articles = []
    seen = set()

    for name, url in FEEDS:
        print(f"📡 {name} ...")
        try:
            feed = feedparser.parse(url)
        except Exception as e:
            print(f"   ⚠️  Fehler: {e}")
            continue

        count = 0
        for entry in feed.entries:
            link = getattr(entry, "link", "")
            if not link or link in seen:
                continue

            pub = parse_date(entry)
            if pub and pub < cutoff:
                continue

            title = getattr(entry, "title", "")
            rss_summary = clean_html(getattr(entry, "summary", ""))

            if not matches_any(f"{title} {rss_summary}", ALL_KEYWORDS):
                continue

            seen.add(link)
            articles.append({
                "title": title,
                "source": name,
                "url": link,
                "rss_summary": rss_summary[:300],
                "date": pub or datetime.now(timezone.utc),
                "_entry": entry,  # Feed-Entry für RSS-Volltext-Extraktion
            })
            count += 1

        print(f"   ✅ {count} Artikel")

    articles.sort(key=lambda a: a["date"], reverse=True)
    if len(articles) > MAX_ARTICLES:
        print(f"\n✂️  Gekürzt auf {MAX_ARTICLES} Artikel")
        articles = articles[:MAX_ARTICLES]

    print(f"\n📊 {len(articles)} Artikel gefunden, hole jetzt Volltexte ...\n")

    # Volltexte holen und zusammenfassen
    paywall_count = 0
    for i, a in enumerate(articles):
        print(f"   📄 [{i+1}/{len(articles)}] {a['title'][:60]}...")
        full_text, text_source = get_best_text(a["_entry"], a["url"])
        a["full_text"] = full_text

        if full_text:
            summary = summarize_text(full_text, a["title"], num_sentences=3)
            a["summary"] = summary or a["rss_summary"]
            a["word_count"] = len(full_text.split())
            source_label = {"web": "Volltext", "rss": "RSS-Inhalt",
                           "partial": "Teiltext"}[text_source]
            print(f"      ✅ {a['word_count']} Wörter ({source_label}) → Zusammenfassung erstellt")
            if text_source != "web":
                paywall_count += 1
        else:
            a["summary"] = a["rss_summary"]
            a["word_count"] = 0
            paywall_count += 1
            print(f"      ℹ️  Kein Volltext verfügbar, nutze RSS-Vorschau")

        del a["_entry"]  # Entry nicht mehr nötig

        a["category"] = categorize(f"{a['title']} {a['summary']}")

    return articles


# ── Word-Dokument erstellen ────────────────────────────────

def add_hyperlink(paragraph, text, url):
    """Fügt einen klickbaren Hyperlink in einen Absatz ein."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})
    color = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "2E75B6"})
    u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    sz = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): "20"})
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    new_run.append(rPr)
    t = paragraph._element.makeelement(qn("w:t"), {})
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def build_word_doc(articles, filename):
    """Erstellt ein formatiertes Word-Dokument mit dem HR-Digest."""
    today = datetime.now()
    kw = today.isocalendar()[1]
    doc = DocxDocument()

    # ── Seitenränder ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles anpassen ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Titel ──
    title = doc.add_heading("HR News Digest", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Kalenderwoche {kw} | {today.strftime('%d.%m.%Y')} | "
                      f"{len(articles)} Artikel aus {len(FEEDS)} Quellen")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Trennlinie ──
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    if not articles:
        doc.add_paragraph("Keine relevanten HR-News diese Woche gefunden.")
        doc.save(filename)
        return

    # ── Top 5 ──
    doc.add_heading("Die 5 wichtigsten Themen", level=1)

    for i, a in enumerate(articles[:5], 1):
        p = doc.add_paragraph()
        num_run = p.add_run(f"{i}. ")
        num_run.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

        title_run = p.add_run(a["title"])
        title_run.bold = True
        title_run.font.size = Pt(11)

        source_run = p.add_run(f"  ({a['source']})")
        source_run.font.size = Pt(9)
        source_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        if a.get("summary"):
            summary_p = doc.add_paragraph()
            summary_run = summary_p.add_run(a["summary"][:300])
            summary_run.font.size = Pt(9)
            summary_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            summary_p.paragraph_format.left_indent = Cm(0.5)
            summary_p.paragraph_format.space_after = Pt(8)

    # ── Trennlinie ──
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Nach Kategorie ──
    grouped = defaultdict(list)
    for a in articles:
        grouped[a["category"]].append(a)

    for cat in list(CATEGORIES.keys()) + ["📌 Sonstiges"]:
        if cat not in grouped:
            continue
        items = grouped[cat]

        h = doc.add_heading(cat, level=1)
        count_p = doc.add_paragraph()
        count_run = count_p.add_run(f"{len(items)} Artikel")
        count_run.font.size = Pt(9)
        count_run.font.italic = True
        count_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        for a in items:
            date_str = a["date"].strftime("%d.%m.")

            # Artikel-Titel
            art_h = doc.add_heading(a["title"], level=2)
            for run in art_h.runs:
                run.font.size = Pt(12)

            # Quelle + Datum + Link
            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run(f"{a['source']} – {date_str}  |  ")
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            add_hyperlink(meta_p, "Zum Artikel →", a["url"])

            # Zusammenfassung
            if a.get("summary"):
                sum_p = doc.add_paragraph()
                sum_text = a["summary"]
                # Paywall-Hinweis
                if a.get("word_count", 0) < 80 and is_paywall_domain(a["url"]):
                    sum_text += "\n\n🔒 Volltext nur mit Abo verfügbar"
                sum_run = sum_p.add_run(sum_text)
                sum_run.font.size = Pt(10)
                sum_p.paragraph_format.space_after = Pt(10)

    # ── Statistik-Tabelle ──
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    doc.add_heading("Statistik", level=1)

    total_words = sum(a.get("word_count", 0) for a in articles)
    with_text = sum(1 for a in articles if a.get("word_count", 0) > 0)

    stats = [
        ("Quellen durchsucht", str(len(FEEDS))),
        ("Relevante Artikel", str(len(articles))),
        ("Wörter verarbeitet", f"{total_words:,}".replace(",", ".")),
        ("Volltexte extrahiert", f"{with_text}/{len(articles)}"),
    ]
    for cat in list(CATEGORIES.keys()) + ["📌 Sonstiges"]:
        if cat in grouped:
            stats.append((cat, f"{len(grouped[cat])} Artikel"))

    table = doc.add_table(rows=len(stats), cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(stats):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # ── Speichern ──
    doc.save(filename)


# ── Main ───────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 55)
    print("  🚀 HR News Digest – Word-Dokument")
    print("=" * 55 + "\n")

    articles = scrape()

    # Word speichern
    today = datetime.now()
    kw = today.isocalendar()[1]
    filename = f"hr-digest-{today.strftime('%Y-%m-%d')}-kw{kw}.docx"

    build_word_doc(articles, filename)

    print(f"\n{'=' * 55}")
    print(f"✅ Gespeichert: {filename}")
    print(f"📄 {len(articles)} Artikel in Word-Dokument")
    print(f"{'=' * 55}")
    print(f"\nDu kannst die Datei jetzt mit Word öffnen: {filename}")
